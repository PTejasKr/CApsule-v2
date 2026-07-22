from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TARGET = Path(r"C:\Users\punya\Desktop\Capsule-Report pre edit .docx")
BACKUP = Path(r"C:\Users\punya\Desktop\capsule\Capsule-Report pre edit backup.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C2CC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def apply_font(run, name="Calibri", size=11, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, before=0, after=6, line=1.12, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def clear_body_after_paragraph(doc, keep_paragraph_index):
    body = doc._body._element
    paragraph_count = -1
    cut_index = None
    children = list(body)
    for i, child in enumerate(children):
        if child.tag == qn("w:p"):
            paragraph_count += 1
            if paragraph_count == keep_paragraph_index:
                cut_index = i + 1
                break
    if cut_index is None:
        raise RuntimeError("Could not find acknowledgement boundary paragraph.")
    for child in children[cut_index:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def make_styles(doc):
    styles = doc.styles
    for missing in ("List Bullet", "List Number"):
        try:
            styles[missing]
        except KeyError:
            styles.add_style(missing, WD_STYLE_TYPE.PARAGRAPH)
    for name, size, color, bold in (
        ("Capsule Body", 11, None, False),
        ("Capsule Heading 1", 16, "1F4D78", True),
        ("Capsule Heading 2", 13, "2E74B5", True),
        ("Capsule Heading 3", 12, "1F4D78", True),
    ):
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = bold
        if color:
            style.font.color.rgb = RGBColor.from_string(color)


def add_h(doc, text, level=1):
    p = doc.add_paragraph(style=f"Capsule Heading {level}")
    r = p.add_run(text)
    apply_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 12), color="1F4D78" if level == 3 else "2E74B5", bold=True)
    format_paragraph(p, before=14 if level == 1 else 10, after=6)
    return p


def add_p(doc, text="", bold_prefix=None):
    p = doc.add_paragraph(style="Capsule Body")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        apply_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        apply_font(r2)
    else:
        r = p.add_run(text)
        apply_font(r)
    format_paragraph(p)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        apply_font(r)
        format_paragraph(p, after=4, line=1.12)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        apply_font(r)
        format_paragraph(p, after=4, line=1.12)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, 9360)
    set_table_borders(table, color="D8DEE6", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FA")
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    p = cell.paragraphs[0]
    for line in text.strip("\n").split("\n"):
        r = p.add_run(line.rstrip())
        apply_font(r, name="Consolas", size=9)
        p.add_run("\n")
    format_paragraph(p, after=0, line=1.0)
    doc.add_paragraph()


def add_image_placeholder(doc, caption):
    p = doc.add_paragraph()
    r = p.add_run(caption)
    apply_font(r, size=10, color="555555", italic=True)
    format_paragraph(p, before=8, after=3)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, 9360)
    set_table_borders(table, color="9AA8B5", size="8")
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=520, bottom=520, start=160, end=160)
    set_cell_shading(cell, "FAFBFC")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = para.add_run("[image]")
    apply_font(rr, size=14, color="666666", bold=True)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_width(table, 9360)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        apply_font(r, size=10, color="1F3A5F", bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(text)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            apply_font(r, size=9.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_entity_schema(doc, name, rows):
    add_h(doc, name, 3)
    add_table(doc, ["Field", "Type", "Constraints", "Purpose"], rows, widths=[1.45, 1.25, 1.75, 2.1])


def build_report(doc):
    add_h(doc, "TABLE OF CONTENTS", 1)
    toc = [
        "List of Tables",
        "List of Figures",
        "List of Abbreviations",
        "Abstract",
        "Introduction and Problem Statement",
        "Industry Overview",
        "Work Responsibility during Internship",
        "Literature Survey",
        "Design Methodology",
        "System Architecture",
        "Database Schema and ER Design",
        "Implementation Details",
        "Testing, Security, and Deployment",
        "Results",
        "Conclusion",
        "Summary",
        "References",
        "Bibliography",
        "Annexure - Interim Report",
    ]
    for idx, item in enumerate(toc, start=1):
        add_p(doc, f"{idx}. {item}")

    add_h(doc, "1. LIST OF TABLES", 1)
    for item in [
        "Table 1: Capsule technology stack",
        "Table 2: Functional and non-functional requirements",
        "Table 3: API endpoint catalogue",
        "Table 4: Logical database schema overview",
        "Table 5: Pydantic request and response schemas",
        "Table 6: Testing strategy and quality gates",
        "Table 7: Deployment and operations checklist",
    ]:
        add_p(doc, item)

    add_h(doc, "2. LIST OF FIGURES", 1)
    for item in [
        "Figure 1: Capsule high-level architecture",
        "Figure 2: Pull request analysis sequence",
        "Figure 3: Database ER diagram",
        "Figure 4: Chrome extension user flow",
        "Figure 5: CI/CD deployment topology",
    ]:
        add_p(doc, item)

    add_h(doc, "3. LIST OF ABBREVIATIONS", 1)
    add_table(doc, ["Abbreviation", "Meaning"], [
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("BRD", "Business Requirements Document"),
        ("CI/CD", "Continuous Integration and Continuous Delivery"),
        ("ER", "Entity Relationship"),
        ("HMAC", "Hash-based Message Authentication Code"),
        ("LLM", "Large Language Model"),
        ("PR", "Pull Request"),
        ("RBAC", "Role-Based Access Control"),
        ("RAG", "Retrieval-Augmented Generation"),
    ], widths=[1.7, 4.8])

    add_h(doc, "4. ABSTRACT", 1)
    add_p(doc, "Capsule is an AI-powered CI/CD companion designed to analyze GitHub pull requests against a Business Requirements Document, detect workflow impact, publish structured PR summaries, and support automated changelog generation. The system combines a FastAPI backend, asynchronous task processing, persistent PR analysis storage, GitHub and Jenkins integration, and a Chrome extension interface that surfaces analysis results directly on GitHub pull request pages.")
    add_p(doc, "The project addresses a common engineering governance problem: teams often merge code without a consistent link between technical changes, business requirements, workflow impact, and release documentation. Capsule reduces this gap by retrieving PR metadata and diffs, grounding AI analysis in BRD content, validating generated findings against the actual changed files, and storing confidence-scored results for approval, repair, preview, and changelog workflows.")
    add_image_placeholder(doc, "Figure 1: Suggested screenshot area for Capsule dashboard or extension summary panel.")
    add_p(doc, "Keywords: Capsule, Pull Request Analysis, Business Requirements, FastAPI, Chrome Extension, GitHub Webhooks, Jenkins, Celery, Redis, PostgreSQL, SQLite, AI Governance, Changelog Automation.")

    add_h(doc, "5. INTRODUCTION AND PROBLEM STATEMENT", 1)
    add_h(doc, "5.1 Introduction", 2)
    add_p(doc, "Modern software teams rely on pull requests as the primary checkpoint before production changes are merged. Although PR reviews are effective for code quality, they can miss whether a change violates business rules, modifies a critical workflow, or requires a release note. Capsule introduces an automated analysis layer that turns PR events into traceable, reviewable engineering intelligence.")
    add_p(doc, "The application is built around a backend service that receives GitHub and Jenkins webhooks, fetches PR details and diffs, loads the active BRD, asks an AI engine to produce structured findings, validates those findings, persists them, and exposes them through API endpoints. A Chrome extension then lets engineers and approvers inspect the summary without leaving GitHub.")
    add_h(doc, "5.2 Problem Statement", 2)
    add_bullets(doc, [
        "Manual PR review does not consistently compare changes against documented business requirements.",
        "Release notes and changelogs are often written late, manually, and with incomplete technical context.",
        "Large diffs are difficult to analyze as a whole, especially when workflow changes span multiple files.",
        "AI-generated summaries can hallucinate file names or impacts unless grounded and validated.",
        "Teams need a lightweight approval and repair workflow that fits existing GitHub and CI/CD usage.",
    ])
    add_h(doc, "5.3 Objectives", 2)
    add_numbered(doc, [
        "Build a PR analysis backend using FastAPI with authenticated API access.",
        "Integrate GitHub webhooks, Jenkins triggers, and background task processing.",
        "Represent AI output through strict Pydantic schemas for predictable storage and API responses.",
        "Persist profiles, BRD versions, repository mappings, PR analyses, changelog entries, and audit logs.",
        "Expose a Chrome extension panel for GitHub PR pages and administrative approval workflows.",
        "Provide ER diagrams, schema design, deployment topology, and testing strategy for maintainability.",
    ])
    add_h(doc, "5.4 Scope", 2)
    add_p(doc, "The scope includes PR analysis, BRD comparison, workflow impact detection, analysis approval, summary repair, changelog preview and generation, multi-profile repository mapping, BRD version management, workflow diagram generation, and browser extension integration. Out of scope items include full project management replacement, human reviewer substitution, and guaranteed semantic correctness without human approval for high-risk changes.")

    add_h(doc, "6. INDUSTRY OVERVIEW", 1)
    add_p(doc, "Engineering organizations increasingly use AI-assisted review, compliance automation, and release intelligence to improve delivery speed while preserving governance. Tools such as code scanners, policy engines, and CI/CD quality gates are mature, but business-rule-aware PR interpretation remains a growing area. Capsule sits in this space by connecting code diffs to BRD-defined business workflows.")
    add_p(doc, "The industry trend is toward event-driven DevOps systems where pull requests, builds, deployments, and release notes are not isolated activities. Capsule follows that pattern by using webhooks, asynchronous workers, persistent audit records, and extension-based user experience.")
    add_table(doc, ["Layer", "Technology", "Reason for Selection"], [
        ("Backend API", "FastAPI", "Async-friendly Python web framework with automatic OpenAPI support."),
        ("Data Models", "Pydantic", "Strict validation for PR summaries, changelog entries, and request payloads."),
        ("Database", "SQLite / PostgreSQL", "Local development simplicity with production-ready relational deployment."),
        ("Async Queue", "Celery", "Background processing for long-running PR analysis and changelog generation."),
        ("Cache/Broker", "Redis", "Fast broker/cache support for worker coordination."),
        ("Frontend", "Chrome Extension MV3", "Direct GitHub PR-page integration without a separate portal requirement."),
        ("Deployment", "Docker Compose, Nginx, Helm", "Containerized local and production-ready deployment paths."),
    ], widths=[1.25, 1.65, 3.6])

    add_h(doc, "7. WORK RESPONSIBILITY DURING INTERNSHIP", 1)
    add_p(doc, "During the internship period, the work centered on understanding the Capsule architecture, analyzing the existing codebase, documenting the application flow, mapping schemas and relationships, and converting implementation details into a formal engineering report.")
    add_h(doc, "7.1 Key Responsibilities", 2)
    add_bullets(doc, [
        "Studied FastAPI routing, middleware, service classes, and Pydantic models.",
        "Mapped the database tables created for profiles, teams, BRD versions, PR analysis, changelogs, and auditing.",
        "Documented the GitHub webhook, Jenkins trigger, and Chrome extension flows.",
        "Prepared logical schema tables and ER relationships for future maintenance.",
        "Identified deployment components including PostgreSQL, Redis, Celery worker, API server, Nginx, Docker, and Helm.",
    ])
    add_image_placeholder(doc, "Figure 2: Suggested image area for internship work timeline or module ownership chart.")

    add_h(doc, "8. LITERATURE SURVEY", 1)
    add_h(doc, "8.1 Pull Request Review and CI/CD Governance", 2)
    add_p(doc, "Pull request review is a widely adopted practice for peer validation before code integration. CI/CD governance extends this practice through automated checks, status checks, build pipelines, and policy enforcement. Capsule enhances this workflow by adding BRD-aware analysis and release documentation support.")
    add_h(doc, "8.2 AI-Assisted Code Analysis", 2)
    add_p(doc, "Large language models can summarize diffs, classify changes, and explain workflow impact. However, their output must be constrained by schemas and validated against actual changed files. Capsule implements this through JSON response requirements, confidence scores, file-path cross-validation, a critic pass, and audit logging.")
    add_h(doc, "8.3 Browser Extension Interfaces", 2)
    add_p(doc, "Browser extensions provide contextual user interfaces inside existing tools. Capsule uses Manifest V3, a background service worker, popup/options pages, and content scripts targeting GitHub PR URLs. This keeps analysis access close to the developer's natural review surface.")
    add_h(doc, "8.4 Relational Schema Design", 2)
    add_p(doc, "Because Capsule manages profiles, BRD versions, teams, repository mappings, PR results, changelog entries, and audit logs, a relational schema is appropriate. Primary keys, foreign keys, and unique constraints help maintain consistency between organizational settings and PR-level output.")

    add_h(doc, "9. DESIGN METHODOLOGY", 1)
    add_h(doc, "9.1 Requirement Classification", 2)
    add_table(doc, ["Requirement", "Type", "Implementation"], [
        ("Analyze PR diff", "Functional", "GitHubService fetches PR details and diff; AIEngine performs map-reduce analysis."),
        ("Compare against BRD", "Functional", "BRDManager loads active profile BRD and passes it into AI prompts."),
        ("Store analysis", "Functional", "pr_analyses table stores summary, branch, approval, changes JSON, workflow JSON, and confidence."),
        ("Approve or reject", "Functional", "Admin/API routes update approval state or remove pending analyses."),
        ("Generate changelog", "Functional", "ChangelogService converts approved PRSummary into a versioned entry and pushes it."),
        ("Prevent fake AI references", "Non-functional", "AIEngine validates file paths against actual diff files and applies confidence penalties."),
        ("Scale heavy processing", "Non-functional", "Celery queues analyze_pr_task and generate_changelog_task in production."),
        ("Secure API access", "Non-functional", "API key middleware and GitHub HMAC signature verification protect external entry points."),
    ], widths=[1.7, 1.1, 3.7])
    add_h(doc, "9.2 Architectural Principles", 2)
    add_bullets(doc, [
        "Event-driven processing: GitHub and Jenkins events initiate analysis without manual polling.",
        "Schema-first contracts: Pydantic models define API payloads and AI output shape.",
        "Separation of concerns: routers handle HTTP, services handle integrations, and database helpers centralize persistence.",
        "Defense in depth: API key checks, GitHub signature verification, repository name validation, and AI output validation reduce risk.",
        "Deployment portability: SQLite supports local use while PostgreSQL, Redis, Celery, Nginx, Docker, and Helm support production environments.",
    ])

    add_h(doc, "10. SYSTEM ARCHITECTURE", 1)
    add_h(doc, "10.1 High-Level Architecture", 2)
    add_code_block(doc, """
graph TD
    Developer[Developer opens or updates PR]
    GitHub[GitHub Pull Request]
    Jenkins[Jenkins Pipeline]
    Extension[Chrome Extension]
    API[Capsule FastAPI Backend]
    Worker[Celery Worker]
    BRD[BRD Manager]
    AI[AI Engine and Provider Router]
    DB[(SQLite or PostgreSQL)]
    Redis[(Redis Broker/Cache)]
    Changelog[Release Repository / Changelog Branch]

    Developer --> GitHub
    GitHub --> API
    Jenkins --> API
    API --> Redis
    Redis --> Worker
    Worker --> BRD
    Worker --> AI
    Worker --> DB
    Extension --> API
    API --> DB
    API --> Changelog
""")
    add_image_placeholder(doc, "Figure 3: Suggested rendered architecture diagram image.")
    add_h(doc, "10.2 Request Processing Flow", 2)
    add_numbered(doc, [
        "A pull_request event reaches /webhooks/github or Jenkins calls /webhooks/jenkins.",
        "The request is authenticated using GitHub HMAC or X-API-Key validation.",
        "In production, Celery receives an asynchronous task; in tests or mock mode, processing can run synchronously.",
        "GitHubService fetches PR metadata, changed files, and unified diff text.",
        "BRDManager loads the active BRD for the selected profile.",
        "AIEngine chunks the diff, analyzes each chunk, optionally performs a global reduce pass, and runs critic validation.",
        "Validated PRSummary data is persisted in pr_analyses and optionally posted back to the PR as a comment.",
        "Approved or merged PRs can generate changelog entries and push them to the configured release repository.",
    ])
    add_h(doc, "10.3 API Endpoint Catalogue", 2)
    add_table(doc, ["Endpoint", "Method", "Purpose"], [
        ("/api/health", "GET", "Health check for the backend service."),
        ("/webhooks/github", "POST", "Receives GitHub pull_request events with signature verification."),
        ("/webhooks/jenkins", "POST", "Allows Jenkins to trigger PR analysis through API-key authentication."),
        ("/webhooks/task/{task_id}", "GET", "Returns Celery task state for queued work."),
        ("/api/pr/{pr_number}/summary", "GET", "Returns an approved or main-branch PR summary."),
        ("/api/pr/{pr_number}/workflow-impact", "GET", "Returns workflow impact only."),
        ("/api/pr/{pr_number}/changelog-preview", "GET", "Generates a changelog preview from stored analysis."),
        ("/api/pr/{pr_number}/approve", "POST", "Approves a PR analysis and attempts changelog generation."),
        ("/api/pr/{pr_number}/repair", "POST", "Saves a manually repaired summary."),
        ("/api/pr/{pr_number}/auto-repair", "POST", "Uses the multi-agent loop to propose and commit repairs."),
        ("/api/profiles/", "GET/POST", "Lists or creates configuration profiles."),
        ("/api/profiles/{id}/brd/upload", "POST", "Uploads a BRD version for a profile."),
        ("/api/workflow/diagram", "POST", "Generates Mermaid workflow text and a QuickChart render URL."),
    ], widths=[2.15, 0.8, 3.55])

    add_h(doc, "11. DATABASE SCHEMA AND ER DESIGN", 1)
    add_h(doc, "11.1 Logical Schema Overview", 2)
    add_table(doc, ["Table", "Primary Key", "Important Fields", "Purpose"], [
        ("profiles", "id", "name, changelog_repo, ai_model, brd_content, github_token, custom_rules, is_super_admin", "Stores analysis profile and integration settings."),
        ("brd_versions", "id", "content, version, uploaded_at, hash, profile_id", "Maintains BRD history per profile."),
        ("teams", "id", "name, created_by, created_at", "Groups users/profiles for shared repository ownership."),
        ("team_members", "team_id + profile_id", "role", "Maps profiles into teams with a role."),
        ("repository_mappings", "source_repo", "team_id, profile_id, created_at", "Maps repositories to analysis profiles or teams."),
        ("pr_analyses", "pr_number + repo", "title, summary, branch, approved, changes_json, workflow_impact_json, confidence_score", "Stores AI analysis results for each pull request."),
        ("changelog_entries", "id", "version, date, technical_changes_json, workflow_changes_json, lines_added, lines_deleted, pr_number", "Stores generated release log entries."),
        ("audit_log", "id", "pr_number, input_hash, output_json, model, tokens, latency_ms, timestamp", "Records AI transaction metadata for traceability."),
    ], widths=[1.3, 1.15, 2.6, 1.45])

    add_h(doc, "11.2 ER Diagram", 2)
    add_code_block(doc, """
erDiagram
    PROFILES ||--o{ BRD_VERSIONS : owns
    PROFILES ||--o{ TEAMS : creates
    PROFILES ||--o{ TEAM_MEMBERS : belongs_to
    TEAMS ||--o{ TEAM_MEMBERS : contains
    PROFILES ||--o{ REPOSITORY_MAPPINGS : configures
    TEAMS ||--o{ REPOSITORY_MAPPINGS : owns
    REPOSITORY_MAPPINGS ||--o{ PR_ANALYSES : routes
    PR_ANALYSES ||--o{ CHANGELOG_ENTRIES : produces
    PR_ANALYSES ||--o{ AUDIT_LOG : records
""")
    add_image_placeholder(doc, "Figure 4: Suggested rendered ER diagram image.")

    add_h(doc, "11.3 Detailed Table Schemas", 2)
    add_entity_schema(doc, "profiles", [
        ("id", "INTEGER/SERIAL", "Primary key", "Unique profile identifier."),
        ("name", "TEXT", "NOT NULL, UNIQUE", "Profile display name."),
        ("changelog_repo", "TEXT", "NOT NULL", "Target owner/repo for changelog output."),
        ("ai_model", "TEXT", "NOT NULL", "Preferred model identifier."),
        ("brd_content", "TEXT", "Nullable", "Profile-specific BRD fallback text."),
        ("github_token", "TEXT", "Nullable", "Token used for one-click repository setup."),
        ("custom_rules", "TEXT", "Nullable", "Additional analysis policy text."),
        ("is_super_admin", "BOOLEAN", "Default false", "Marks the default administrative profile."),
        ("created_at", "TIMESTAMP", "Default current timestamp", "Creation time."),
    ])
    add_entity_schema(doc, "brd_versions", [
        ("id", "INTEGER/SERIAL", "Primary key", "Unique BRD version row."),
        ("content", "TEXT", "NOT NULL", "Uploaded BRD body."),
        ("version", "TEXT", "NOT NULL", "Human-readable BRD version."),
        ("uploaded_at", "TIMESTAMP", "Default current timestamp", "Upload time."),
        ("hash", "TEXT", "NOT NULL, UNIQUE", "Content hash for duplicate detection."),
        ("profile_id", "INTEGER", "FK profiles(id), cascade delete", "Owner profile."),
    ])
    add_entity_schema(doc, "repository_mappings", [
        ("source_repo", "TEXT", "Primary key", "GitHub owner/repo source identifier."),
        ("team_id", "INTEGER", "FK teams(id), nullable", "Team that owns this repository."),
        ("profile_id", "INTEGER", "FK profiles(id), nullable", "Profile used for analysis settings."),
        ("created_at", "TIMESTAMP", "Default current timestamp", "Mapping creation time."),
    ])
    add_entity_schema(doc, "pr_analyses", [
        ("pr_number", "INTEGER", "Composite primary key", "Pull request number."),
        ("repo", "TEXT", "Composite primary key", "Repository owner/name."),
        ("title", "TEXT", "Nullable", "PR title."),
        ("summary", "TEXT", "Nullable", "Current human-facing AI summary."),
        ("original_summary", "TEXT", "Nullable", "Unedited AI summary for comparison."),
        ("brd_comparison", "TEXT", "Nullable", "Explanation of BRD alignment or violation."),
        ("branch", "TEXT", "Nullable", "Source or target branch used for access rules."),
        ("approved", "BOOLEAN", "Default false", "Controls visibility and changelog generation."),
        ("changes_json", "TEXT", "JSON payload", "Serialized list of ChangeItem objects."),
        ("workflow_impact_json", "TEXT", "JSON payload", "Serialized WorkflowImpact object."),
        ("confidence_score", "REAL", "0.0 to 1.0", "Overall analysis confidence."),
        ("analyzed_at", "TIMESTAMP", "Default current timestamp", "Analysis time."),
        ("author", "TEXT", "Nullable", "GitHub author or merger metadata."),
        ("merged_at", "TEXT", "Nullable", "Merge timestamp if available."),
    ])
    add_entity_schema(doc, "changelog_entries", [
        ("id", "INTEGER/SERIAL", "Primary key", "Unique changelog entry."),
        ("version", "TEXT", "NOT NULL", "Semantic version produced by changelog logic."),
        ("date", "TEXT", "NOT NULL", "Release date."),
        ("technical_changes_json", "TEXT", "JSON list", "Technical release note lines."),
        ("workflow_changes_json", "TEXT", "JSON list", "Workflow impact release note lines."),
        ("lines_added", "INTEGER", "Default 0", "Added line count."),
        ("lines_deleted", "INTEGER", "Default 0", "Deleted line count."),
        ("pr_number", "INTEGER", "Nullable", "Associated PR number."),
        ("pushed_at", "TIMESTAMP", "Default current timestamp", "Push or creation time."),
    ])
    add_entity_schema(doc, "audit_log", [
        ("id", "INTEGER/SERIAL", "Primary key", "Unique audit row."),
        ("pr_number", "INTEGER", "Nullable", "Associated PR number."),
        ("input_hash", "TEXT", "Nullable", "Hash of diff plus BRD input."),
        ("output_json", "TEXT", "Nullable", "Stored AI output for traceability."),
        ("model", "TEXT", "Nullable", "Provider/model used."),
        ("tokens", "INTEGER", "Nullable", "Token count if collected."),
        ("latency_ms", "REAL", "Nullable", "Analysis latency."),
        ("timestamp", "TIMESTAMP", "Default current timestamp", "Audit event time."),
    ])

    add_h(doc, "11.4 Pydantic Data Schemas", 2)
    add_table(doc, ["Schema", "Important Fields", "Use"], [
        ("ChangeItem", "file, line_range, change_type, description, confidence, reasoning_trace", "Represents a single changed file or code region."),
        ("WorkflowImpact", "has_impact, severity, impact_description, affected_workflows, before_state, after_state", "Captures business workflow effect."),
        ("PRSummary", "pr_number, repo, branch, title, summary, brd_comparison, changes, workflow_impact, confidence_score", "Main analysis response returned to users."),
        ("ChangelogEntry", "version, date, technical_changes, workflow_changes, lines_added, lines_deleted, pr_number", "Release note/changelog unit."),
        ("ProfileCreate", "name, changelog_repo, ai_model, brd_content, github_token, custom_rules", "Profile creation and update request."),
        ("RepositoryMappingCreate", "source_repo, profile_id, team_id", "Binds a GitHub repo to a Capsule profile or team."),
        ("BRDUploadResponse", "status, version, hash, uploaded_at", "Confirms uploaded BRD version metadata."),
    ], widths=[1.45, 3.15, 1.9])

    add_h(doc, "12. IMPLEMENTATION DETAILS", 1)
    add_h(doc, "12.1 Backend Implementation", 2)
    add_p(doc, "The backend is initialized in extension/backend/main.py. It creates the FastAPI application, configures CORS for extension and API requests, initializes the database during lifespan startup, loads the default BRD, and mounts routers for authentication, administration, webhooks, API actions, profiles, and teams.")
    add_p(doc, "The router layer stays thin: it validates requests, applies authentication dependencies, and delegates work to service classes. The service layer contains GitHub integration, AI analysis, changelog generation, BRD management, routing across AI providers, guardrails, and PR orchestration.")
    add_h(doc, "12.2 AI Analysis Engine", 2)
    add_numbered(doc, [
        "Parse changed file paths from the unified diff.",
        "Chunk the diff at file-aware boundaries to stay within model context limits.",
        "Analyze chunks concurrently with a BRD-grounded system prompt.",
        "Merge chunk findings and optionally run a global reduce pass.",
        "Run a critic review against the raw diff.",
        "Cross-validate every returned file path against actual changed files.",
        "Apply confidence penalties when fabricated references are removed.",
        "Persist final structured PRSummary and audit metadata.",
    ])
    add_h(doc, "12.3 Chrome Extension", 2)
    add_p(doc, "The Chrome extension uses Manifest V3. Content scripts are injected on GitHub pull request pages, while popup and options pages provide summary access and configuration. The extension authenticates with the backend using the configured API key and is intentionally lightweight, relying on the backend for analysis and persistence.")
    add_image_placeholder(doc, "Figure 5: Suggested Chrome extension popup screenshot.")
    add_h(doc, "12.4 Admin and Approval Workflow", 2)
    add_p(doc, "Capsule supports a practical approval flow. Unapproved feature-branch analyses are protected from normal retrieval. Admin routes can approve, reject, repair, compare, or auto-repair a summary. Approval can also trigger changelog generation, allowing organizations to keep human oversight in the loop while automating repetitive release documentation.")

    add_h(doc, "13. TESTING, SECURITY, AND DEPLOYMENT", 1)
    add_h(doc, "13.1 Testing Strategy", 2)
    add_table(doc, ["Area", "Tests/Checks", "Expected Assurance"], [
        ("Unit tests", "AI engine, async behavior, schema reconstruction", "Core logic works without external services."),
        ("Integration tests", "Webhooks, approval workflow, incidents, pipeline approval", "HTTP routes and persistence paths behave together."),
        ("Security tests", "Signature verification, data leakage, AI integrity", "External inputs and generated outputs are guarded."),
        ("Mock mode", "Sandbox webhook events", "Demo and test flows run without real GitHub calls."),
        ("Load testing", "locustfile.py", "Basic performance profile under repeated API calls."),
    ], widths=[1.4, 2.55, 2.55])
    add_h(doc, "13.2 Security Controls", 2)
    add_bullets(doc, [
        "X-API-Key validation protects extension and administrative endpoints.",
        "GitHub webhook requests are verified using the x-hub-signature-256 HMAC flow.",
        "Repository names are validated before being used in API operations.",
        "Production mode disables public FastAPI documentation endpoints.",
        "AI output is cross-validated against actual diff file paths before persistence.",
        "Docker services use no-new-privileges and resource limits in compose configuration.",
    ])
    add_h(doc, "13.3 Deployment Architecture", 2)
    add_p(doc, "The default containerized deployment includes PostgreSQL, Redis, the Capsule API server, a Celery worker, and Nginx. Docker Compose supports local or VM deployment, while Helm chart files provide a Kubernetes-oriented packaging path.")
    add_code_block(doc, """
flowchart LR
    Internet[GitHub/Jenkins/Extension]
    Nginx[Nginx Reverse Proxy]
    API[Capsule API Container]
    Worker[Celery Worker Container]
    Postgres[(PostgreSQL)]
    Redis[(Redis)]
    Volume[(BRD and Data Volumes)]

    Internet --> Nginx
    Nginx --> API
    API --> Redis
    Redis --> Worker
    API --> Postgres
    Worker --> Postgres
    API --> Volume
    Worker --> Volume
""")
    add_image_placeholder(doc, "Figure 6: Suggested deployment topology diagram image.")
    add_table(doc, ["Component", "Port/Interface", "Operational Note"], [
        ("capsule-api", "8000", "FastAPI app receives extension, webhook, and admin requests."),
        ("celery-worker", "Redis broker", "Processes queued analysis and changelog jobs."),
        ("postgres", "5432", "Production relational database."),
        ("redis", "6379", "Broker/cache for background task processing."),
        ("nginx", "80", "Reverse proxy in front of the API service."),
        ("Chrome extension", "GitHub pages + backend URL", "Shows PR summaries in the user's browser."),
    ], widths=[1.55, 1.55, 3.4])

    add_h(doc, "14. RESULTS", 1)
    add_p(doc, "The completed Capsule system provides an end-to-end pipeline from PR event ingestion to structured analysis and release documentation. The design is modular enough for local development, test mocks, Docker-based deployment, and production asynchronous operation.")
    add_bullets(doc, [
        "PR events can be received from GitHub or Jenkins.",
        "AI summaries are grounded in BRD context and persisted in a relational schema.",
        "Workflow impact is represented separately from general technical changes.",
        "Admins can approve, reject, repair, compare, or auto-repair analyses.",
        "Chrome extension users can view analysis in the context of a GitHub pull request.",
        "Changelog entries can be generated and pushed after approval or merge.",
    ])
    add_image_placeholder(doc, "Figure 7: Suggested screenshot area for PR summary result in GitHub.")
    add_image_placeholder(doc, "Figure 8: Suggested screenshot area for generated changelog or weekly changes view.")

    add_h(doc, "15. CONCLUSION", 1)
    add_p(doc, "Capsule demonstrates how AI can be embedded responsibly into the software delivery lifecycle. Instead of only producing free-form summaries, it enforces structured schemas, stores auditable results, validates file references, and keeps human approval available for risky or edited summaries. The result is a practical PR intelligence system that supports developer productivity, release consistency, and business-rule traceability.")
    add_p(doc, "The project also provides a strong foundation for future enhancements such as richer organization-level RBAC, configurable policy packs, deeper CI status enforcement, Slack or Teams notifications, and improved visualization of workflow impact across multiple repositories.")

    add_h(doc, "16. SUMMARY", 1)
    add_p(doc, "Capsule is an AI-powered PR analyzer and changelog automation platform. It includes FastAPI routes, Pydantic schemas, GitHub and Jenkins integrations, AI provider routing, BRD version management, relational persistence, Celery background processing, Redis coordination, Docker/Nginx deployment, and a Chrome extension interface.")
    add_p(doc, "The report body documents Capsule from project motivation through architecture, schema design, ER modeling, implementation, testing, security, deployment, and results. Placeholders marked [image] have been left for future screenshots, rendered diagrams, and interface captures.")

    add_h(doc, "17. REFERENCES", 1)
    for ref in [
        "FastAPI documentation for API routing, dependencies, and application lifespan.",
        "Pydantic documentation for BaseModel validation and schema-driven request/response design.",
        "Chrome Extensions Manifest V3 documentation for service workers, content scripts, permissions, and extension UI.",
        "GitHub Webhooks documentation for pull_request events and HMAC signature verification.",
        "Celery and Redis documentation for asynchronous background task processing.",
        "PostgreSQL and SQLite documentation for relational schema behavior and SQL compatibility.",
        "Docker Compose and Nginx documentation for container orchestration and reverse proxy deployment.",
    ]:
        add_p(doc, ref)

    add_h(doc, "18. BIBLIOGRAPHY", 1)
    for item in [
        "Software Engineering: A Practitioner's Approach - Roger S. Pressman.",
        "Designing Data-Intensive Applications - Martin Kleppmann.",
        "Clean Architecture - Robert C. Martin.",
        "Continuous Delivery - Jez Humble and David Farley.",
        "Building Microservices - Sam Newman.",
    ]:
        add_p(doc, item)

    add_h(doc, "19. ANNEXURE - INTERIM REPORT", 1)
    add_h(doc, "19.1 Future Enhancement Roadmap", 2)
    add_table(doc, ["Enhancement", "Benefit", "Priority"], [
        ("Organization-level RBAC", "Separates admin, reviewer, developer, and auditor responsibilities.", "High"),
        ("Policy pack editor", "Lets teams configure rules without editing source code.", "High"),
        ("Native Slack/Teams alerts", "Sends high-impact PR notifications to collaboration channels.", "Medium"),
        ("Vector BRD retrieval", "Improves BRD grounding for long documents.", "Medium"),
        ("Interactive ER/architecture diagrams", "Improves onboarding for maintainers.", "Medium"),
        ("Deployment observability dashboard", "Tracks queue latency, model latency, and webhook failures.", "Medium"),
    ], widths=[2.05, 3.55, 0.9])
    add_h(doc, "19.2 Image Checklist", 2)
    add_bullets(doc, [
        "Add product logo or project banner on the report body opening page.",
        "Add screenshot of Chrome extension popup on a GitHub PR page.",
        "Add rendered high-level architecture diagram.",
        "Add rendered ER diagram from the Mermaid block.",
        "Add screenshot of API health response or OpenAPI docs in development mode.",
        "Add screenshot of generated changelog or approval dashboard.",
    ])


def main():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    source = BACKUP if BACKUP.exists() else TARGET
    doc = Document(source)
    if not BACKUP.exists():
        backup_doc = Document(TARGET)
        backup_doc.save(BACKUP)
    clear_body_after_paragraph(doc, 105)
    make_styles(doc)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    build_report(doc)
    doc.save(TARGET)


if __name__ == "__main__":
    main()
